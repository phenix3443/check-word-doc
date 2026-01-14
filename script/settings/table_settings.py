#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
表格验证设置模块
提供表格验证相关的通用函数和样式验证
"""

from typing import List, Any, Callable, Tuple, Optional
from dataclasses import dataclass

# 尝试导入文档处理模块
DOCX_AVAILABLE = False
try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    DOCX_AVAILABLE = True
except ImportError:
    pass


@dataclass
class TableStyleSettings:
    """表格样式设置结构体

    表示Word文档中表格的格式设置，包括对齐方式、边框、内边距等属性。
    对应Word表格属性对话框中的设置项。
    """

    # 单元格对齐方式
    cell_horizontal_alignment: str = ""  # 水平对齐：左对齐、居中、右对齐、两端对齐、分散对齐
    cell_vertical_alignment: str = ""  # 垂直对齐：顶端对齐、居中、底端对齐

    # 单元格内边距（磅）
    cell_top_padding: Optional[float] = None  # 上内边距
    cell_bottom_padding: Optional[float] = None  # 下内边距
    cell_left_padding: Optional[float] = None  # 左内边距
    cell_right_padding: Optional[float] = None  # 右内边距

    # 表格对齐方式
    table_alignment: Optional[str] = None  # 表格对齐：左对齐、居中、右对齐

    # 表格尺寸
    table_rows: Optional[int] = None  # 表格行数
    table_columns: Optional[int] = None  # 表格列数

    # 边框设置（可选，如果需要验证边框）
    border_style: Optional[str] = None  # 边框样式

    def __str__(self) -> str:
        """返回表格样式设置的字符串表示"""
        parts = []
        if self.cell_horizontal_alignment:
            parts.append(f"单元格水平对齐：{self.cell_horizontal_alignment}")
        if self.cell_vertical_alignment:
            parts.append(f"单元格垂直对齐：{self.cell_vertical_alignment}")
        if self.table_alignment:
            parts.append(f"表格对齐：{self.table_alignment}")
        if self.table_rows is not None:
            parts.append(f"表格行数：{self.table_rows}")
        if self.table_columns is not None:
            parts.append(f"表格列数：{self.table_columns}")
        if self.cell_top_padding is not None:
            parts.append(f"上内边距：{self.cell_top_padding}磅")
        if self.cell_bottom_padding is not None:
            parts.append(f"下内边距：{self.cell_bottom_padding}磅")
        if self.cell_left_padding is not None:
            parts.append(f"左内边距：{self.cell_left_padding}磅")
        if self.cell_right_padding is not None:
            parts.append(f"右内边距：{self.cell_right_padding}磅")
        if self.border_style:
            parts.append(f"边框样式：{self.border_style}")
        return "；".join(parts) if parts else "无设置"

    def cell_horizontal_alignment_validator(
        self,
    ) -> Optional[Callable[[Any, Any, str, str], Tuple[bool, Optional[str]]]]:
        """创建单元格水平对齐验证器

        Returns:
            验证器函数，符合标准验证器接口 (doc, para_or_cell, field_name, field_value) -> (bool, Optional[str])
            如果未设置对齐方式或 DOCX 不可用，返回 None
        """
        if not DOCX_AVAILABLE:
            return None

        if not self.cell_horizontal_alignment:
            return None

        alignment_mapping = {
            "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
            "居中": WD_ALIGN_PARAGRAPH.CENTER,
            "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
            "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "分散对齐": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
        }

        expected_alignment = alignment_mapping.get(self.cell_horizontal_alignment)
        if expected_alignment is None:
            return None

        expected_text = self.cell_horizontal_alignment

        def validator(
            doc: Any, para_or_cell: Any, field_name: str, field_value: str
        ) -> Tuple[bool, Optional[str]]:
            if para_or_cell is None:
                return True, None

            try:
                # para_or_cell 应该是单元格对象
                if not hasattr(para_or_cell, "_parent"):
                    return True, None  # 不是单元格，跳过检查

                # 获取单元格中的段落对齐方式
                paragraphs = para_or_cell.paragraphs
                if not paragraphs:
                    return True, None

                # 检查所有段落的对齐方式
                for para in paragraphs:
                    if para.alignment != expected_alignment:
                        actual_text = self._get_alignment_text(para.alignment)
                        return (
                            False,
                            f"{field_name}单元格水平对齐应为{expected_text}，当前为{actual_text}",
                        )
            except Exception:
                pass
            return True, None

        return validator

    def cell_vertical_alignment_validator(
        self,
    ) -> Optional[Callable[[Any, Any, str, str], Tuple[bool, Optional[str]]]]:
        """创建单元格垂直对齐验证器

        Returns:
            验证器函数，符合标准验证器接口 (doc, para_or_cell, field_name, field_value) -> (bool, Optional[str])
            如果未设置对齐方式或 DOCX 不可用，返回 None
        """
        if not DOCX_AVAILABLE:
            return None

        if not self.cell_vertical_alignment:
            return None

        alignment_mapping = {
            "顶端对齐": WD_CELL_VERTICAL_ALIGNMENT.TOP,
            "居中": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "底端对齐": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        }

        expected_alignment = alignment_mapping.get(self.cell_vertical_alignment)
        if expected_alignment is None:
            return None

        expected_text = self.cell_vertical_alignment

        def validator(
            doc: Any, para_or_cell: Any, field_name: str, field_value: str
        ) -> Tuple[bool, Optional[str]]:
            if para_or_cell is None:
                return True, None

            try:
                # para_or_cell 应该是单元格对象
                if not hasattr(para_or_cell, "_parent"):
                    return True, None  # 不是单元格，跳过检查

                # 获取单元格的垂直对齐方式
                if para_or_cell.vertical_alignment != expected_alignment:
                    actual_text = self._get_vertical_alignment_text(
                        para_or_cell.vertical_alignment
                    )
                    return (
                        False,
                        f"{field_name}单元格垂直对齐应为{expected_text}，当前为{actual_text}",
                    )
            except Exception:
                pass
            return True, None

        return validator

    def cell_padding_validator(
        self,
    ) -> Optional[Callable[[Any, Any, str, str], Tuple[bool, Optional[str]]]]:
        """创建单元格内边距验证器

        Returns:
            验证器函数，符合标准验证器接口 (doc, para_or_cell, field_name, field_value) -> (bool, Optional[str])
            如果未设置内边距或 DOCX 不可用，返回 None
        """
        if not DOCX_AVAILABLE:
            return None

        # 收集需要检查的内边距
        paddings_to_check = []
        if self.cell_top_padding is not None:
            paddings_to_check.append(("top", "上", self.cell_top_padding))
        if self.cell_bottom_padding is not None:
            paddings_to_check.append(("bottom", "下", self.cell_bottom_padding))
        if self.cell_left_padding is not None:
            paddings_to_check.append(("left", "左", self.cell_left_padding))
        if self.cell_right_padding is not None:
            paddings_to_check.append(("right", "右", self.cell_right_padding))

        if not paddings_to_check:
            return None

        def validator(
            doc: Any, para_or_cell: Any, field_name: str, field_value: str
        ) -> Tuple[bool, Optional[str]]:
            if para_or_cell is None:
                return True, None

            try:
                # para_or_cell 应该是单元格对象
                if not hasattr(para_or_cell, "_parent"):
                    return True, None  # 不是单元格，跳过检查

                # 检查所有需要验证的内边距
                for padding_name, display_name, expected_value in paddings_to_check:
                    # 获取单元格的内边距
                    if padding_name == "top":
                        actual_value = para_or_cell.top_padding.pt if para_or_cell.top_padding else 0
                    elif padding_name == "bottom":
                        actual_value = (
                            para_or_cell.bottom_padding.pt if para_or_cell.bottom_padding else 0
                        )
                    elif padding_name == "left":
                        actual_value = para_or_cell.left_padding.pt if para_or_cell.left_padding else 0
                    elif padding_name == "right":
                        actual_value = (
                            para_or_cell.right_padding.pt if para_or_cell.right_padding else 0
                        )
                    else:
                        continue

                    # 允许小的误差（0.1磅）
                    if abs(actual_value - expected_value) > 0.1:
                        return (
                            False,
                            f"{field_name}单元格{display_name}内边距应为{expected_value}磅，当前为{actual_value}磅",
                        )
            except Exception:
                pass
            return True, None

        return validator

    def table_alignment_validator(
        self,
    ) -> Optional[Callable[[Any, Any, str, str], Tuple[bool, Optional[str]]]]:
        """创建表格对齐验证器

        Returns:
            验证器函数，符合标准验证器接口 (doc, para_or_cell, field_name, field_value) -> (bool, Optional[str])
            如果未设置表格对齐或 DOCX 不可用，返回 None
        """
        if not DOCX_AVAILABLE:
            return None

        if not self.table_alignment:
            return None

        alignment_mapping = {
            "左对齐": WD_TABLE_ALIGNMENT.LEFT,
            "居中": WD_TABLE_ALIGNMENT.CENTER,
            "右对齐": WD_TABLE_ALIGNMENT.RIGHT,
        }

        expected_alignment = alignment_mapping.get(self.table_alignment)
        if expected_alignment is None:
            return None

        expected_text = self.table_alignment

        def validator(
            doc: Any, para_or_cell: Any, field_name: str, field_value: str
        ) -> Tuple[bool, Optional[str]]:
            if para_or_cell is None:
                return True, None

            try:
                # para_or_cell 应该是单元格对象，需要找到其所在的表格
                table = None
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            if cell == para_or_cell:
                                table = t
                                break
                        if table:
                            break
                    if table:
                        break

                if table and table.alignment != expected_alignment:
                    actual_text = self._get_table_alignment_text(table.alignment)
                    return (
                        False,
                        f"{field_name}表格对齐应为{expected_text}，当前为{actual_text}",
                    )
            except Exception:
                pass
            return True, None

        return validator

    def table_dimensions_validator(
        self,
    ) -> Optional[Callable[[Any, Any, str, str], Tuple[bool, Optional[str]]]]:
        """创建表格尺寸验证器（行数和列数）

        Returns:
            验证器函数，符合标准验证器接口 (doc, para_or_cell, field_name, field_value) -> (bool, Optional[str])
            如果未设置行数或列数，返回 None
        """
        # 如果没有设置行数和列数，返回 None
        if self.table_rows is None and self.table_columns is None:
            return None

        expected_rows = self.table_rows
        expected_columns = self.table_columns

        def validator(
            doc: Any, para_or_cell: Any, field_name: str, field_value: str
        ) -> Tuple[bool, Optional[str]]:
            if para_or_cell is None:
                return True, None

            try:
                # para_or_cell 应该是单元格对象，需要找到其所在的表格
                table = None
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            if cell == para_or_cell:
                                table = t
                                break
                        if table:
                            break
                    if table:
                        break

                if not table:
                    return True, None  # 找不到表格，跳过检查

                # 检查行数
                if expected_rows is not None:
                    actual_rows = len(table.rows)
                    if actual_rows != expected_rows:
                        return (
                            False,
                            f"{field_name}表格行数应为{expected_rows}，当前为{actual_rows}",
                        )

                # 检查列数
                if expected_columns is not None:
                    # 获取第一行的列数（假设所有行的列数相同）
                    if len(table.rows) > 0:
                        actual_columns = len(table.rows[0].cells)
                        if actual_columns != expected_columns:
                            return (
                                False,
                                f"{field_name}表格列数应为{expected_columns}，当前为{actual_columns}",
                            )
            except Exception:
                pass
            return True, None

        return validator

    def validators(
        self,
    ) -> List[Callable[[Any, Any, str, str], Tuple[bool, Optional[str]]]]:
        """返回所有相关的验证器列表

        Returns:
            验证器函数列表，每个函数符合标准验证器接口 (doc, para_or_cell, field_name, field_value) -> (bool, Optional[str])
        """
        validators_list = []

        validator = self.cell_horizontal_alignment_validator()
        if validator is not None:
            validators_list.append(validator)

        validator = self.cell_vertical_alignment_validator()
        if validator is not None:
            validators_list.append(validator)

        validator = self.cell_padding_validator()
        if validator is not None:
            validators_list.append(validator)

        validator = self.table_alignment_validator()
        if validator is not None:
            validators_list.append(validator)

        validator = self.table_dimensions_validator()
        if validator is not None:
            validators_list.append(validator)

        return validators_list

    @staticmethod
    def _get_alignment_text(alignment: Any) -> str:
        """获取对齐方式的文本表示"""
        if not DOCX_AVAILABLE:
            return "未知"
        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
            WD_ALIGN_PARAGRAPH.CENTER: "居中",
            WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: "分散对齐",
        }
        return mapping.get(alignment, "未知")

    @staticmethod
    def _get_vertical_alignment_text(alignment: Any) -> str:
        """获取垂直对齐方式的文本表示"""
        if not DOCX_AVAILABLE:
            return "未知"
        mapping = {
            WD_CELL_VERTICAL_ALIGNMENT.TOP: "顶端对齐",
            WD_CELL_VERTICAL_ALIGNMENT.CENTER: "居中",
            WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: "底端对齐",
        }
        return mapping.get(alignment, "未知")

    @staticmethod
    def _get_table_alignment_text(alignment: Any) -> str:
        """获取表格对齐方式的文本表示"""
        if not DOCX_AVAILABLE:
            return "未知"
        mapping = {
            WD_TABLE_ALIGNMENT.LEFT: "左对齐",
            WD_TABLE_ALIGNMENT.CENTER: "居中",
            WD_TABLE_ALIGNMENT.RIGHT: "右对齐",
        }
        return mapping.get(alignment, "未知")


class TableSettings:
    """表格设置类，提供表格验证相关的匹配函数和样式验证"""

    @staticmethod
    def create_caption_match_function(caption_pattern: str = None) -> Callable[[Any, Any], bool]:
        """创建匹配函数，用于判断文档元素是否为表格题注段落

        Args:
            caption_pattern: 题注文本模式（可选），如果提供，则检查文本是否包含该模式

        Returns:
            匹配函数，接受 (doc, para_or_cell) 参数，返回 bool
        """
        def match_func(doc: Any, para_or_cell: Any) -> bool:
            # 必须是段落对象
            if not hasattr(para_or_cell, "runs"):
                return False

            # 获取段落文本
            para_text = para_or_cell.text.strip()
            if not para_text:
                return False

            # 检查是否包含"表"字（表格题注的标识）
            if "表" not in para_text:
                return False

            # 如果提供了模式，检查文本是否匹配
            if caption_pattern:
                if caption_pattern not in para_text:
                    return False

            # 检查段落是否紧邻表格（可选，增强匹配准确性）
            try:
                # 获取文档主体中的所有元素（段落和表格）
                body = doc.element.body
                body_elements = list(body)

                # 找到当前段落在文档中的位置
                para_element = para_or_cell.element
                element_index = body_elements.index(para_element)

                # 检查前后是否有表格
                # 向前查找表格（最多查找前一个元素）
                if element_index > 0:
                    prev_elem = body_elements[element_index - 1]
                    # 检查是否是表格元素（通过检查是否有 rows 属性）
                    if hasattr(prev_elem, "rows") or (
                        hasattr(prev_elem, "tag") and "tbl" in str(prev_elem.tag)
                    ):
                        return True

                # 向后查找表格（最多查找后一个元素）
                if element_index < len(body_elements) - 1:
                    next_elem = body_elements[element_index + 1]
                    # 检查是否是表格元素
                    if hasattr(next_elem, "rows") or (
                        hasattr(next_elem, "tag") and "tbl" in str(next_elem.tag)
                    ):
                        return True
            except (ValueError, AttributeError):
                # 如果无法找到元素位置，至少检查文本内容
                pass

            # 如果文本包含"表"字，即使没有找到相邻表格也返回 True
            return True

        return match_func

    @staticmethod
    def create_table_match_function(keywords: List[str] = None) -> Callable[[Any, Any], bool]:
        """创建匹配函数，用于判断文档元素所在的表格是否为指定的表格

        Args:
            keywords: 关键词列表，用于识别表格（检查表格第一行是否包含这些关键词）

        Returns:
            匹配函数，接受 (doc, para_or_cell) 参数，返回 bool
            注意：para_or_cell 应该是表格单元格（Cell 对象）
        """
        def match_func(doc: Any, para_or_cell: Any) -> bool:
            # 必须是单元格对象
            if not hasattr(para_or_cell, "_parent"):
                return False

            # 通过遍历所有表格来找到包含该单元格的表格
            table = None
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell == para_or_cell:
                            table = t
                            break
                    if table:
                        break
                if table:
                    break

            if not table:
                return False

            # 如果没有提供关键词，则匹配任何表格
            if not keywords:
                return True

            # 检查表格第一行是否包含关键词
            if len(table.rows) == 0:
                return False

            first_row_text = " ".join([cell.text.strip() for cell in table.rows[0].cells])
            return any(keyword in first_row_text for keyword in keywords)

        return match_func

    @staticmethod
    def create_cell_match_function(field_key: str) -> Callable[[Any, Any], bool]:
        """创建匹配函数，用于判断文档元素是否为表格中指定字段的值单元格

        Args:
            field_key: 字段键名（表格第一列的文本），用于匹配对应的值单元格（第二列）

        Returns:
            匹配函数，接受 (doc, para_or_cell) 参数，返回 bool
        """
        def match_func(doc: Any, para_or_cell: Any) -> bool:
            # 必须是单元格对象
            if not hasattr(para_or_cell, "_parent"):
                return False

            # 通过遍历所有表格来找到包含该单元格的行
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) >= 2:
                        key_cell = row.cells[0]
                        value_cell = row.cells[1]
                        key_text = key_cell.text.strip()

                        # 检查键名是否匹配，且当前单元格是值单元格
                        if key_text == field_key and para_or_cell == value_cell:
                            return True

            return False

        return match_func

    @staticmethod
    def create_cell_match_function_with_table(
        keywords: List[str], field_key: str
    ) -> Callable[[Any, Any], bool]:
        """创建匹配函数，用于判断文档元素是否为指定表格中指定字段的值单元格

        Args:
            keywords: 关键词列表，用于识别表格（检查表格第一行是否包含这些关键词）
            field_key: 字段键名（表格第一列的文本），用于匹配对应的值单元格（第二列）

        Returns:
            匹配函数，接受 (doc, para_or_cell) 参数，返回 bool
        """
        table_match = TableSettings.create_table_match_function(keywords)
        cell_match = TableSettings.create_cell_match_function(field_key)

        def match_func(doc: Any, para_or_cell: Any) -> bool:
            return table_match(doc, para_or_cell) and cell_match(doc, para_or_cell)

        return match_func


# 为了向后兼容，提供模块级别的函数作为类的快捷方式
def create_table_caption_match_function(caption_pattern: str = None) -> Callable[[Any, Any], bool]:
    """创建匹配函数，用于判断文档元素是否为表格题注段落（向后兼容函数）"""
    return TableSettings.create_caption_match_function(caption_pattern)


def create_table_match_function(keywords: List[str] = None) -> Callable[[Any, Any], bool]:
    """创建匹配函数，用于判断文档元素所在的表格是否为指定的表格（向后兼容函数）"""
    return TableSettings.create_table_match_function(keywords)


def create_table_cell_match_function(field_key: str) -> Callable[[Any, Any], bool]:
    """创建匹配函数，用于判断文档元素是否为表格中指定字段的值单元格（向后兼容函数）"""
    return TableSettings.create_cell_match_function(field_key)


def create_table_cell_match_function_with_table(
    keywords: List[str], field_key: str
) -> Callable[[Any, Any], bool]:
    """创建匹配函数，用于判断文档元素是否为指定表格中指定字段的值单元格（向后兼容函数）"""
    return TableSettings.create_cell_match_function_with_table(keywords, field_key)
