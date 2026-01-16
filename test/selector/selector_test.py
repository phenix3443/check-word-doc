"""
Selector 测试用例
"""

from script.core.model import ParagraphBlock, TableBlock, Block
from script.core.selector import Selector
from docx import Document


def create_test_blocks():
    """创建测试用的 blocks"""
    doc = Document()
    
    # 创建测试段落
    blocks = []
    
    # 1. 标题
    p1 = doc.add_paragraph("论文标题")
    b1 = ParagraphBlock(index=0, paragraph=p1)
    b1.add_class("title")
    blocks.append(b1)
    
    # 2. 作者列表
    p2 = doc.add_paragraph("张三1*，李四2")
    b2 = ParagraphBlock(index=1, paragraph=p2)
    b2.add_class("author-section")
    b2.add_class("author-list")
    blocks.append(b2)
    
    # 3. 第一个作者单位
    p3 = doc.add_paragraph("1. 北京大学，北京  100871")
    b3 = ParagraphBlock(index=2, paragraph=p3)
    b3.add_class("author-section")
    b3.add_class("author-affiliation")
    blocks.append(b3)
    
    # 4. 第二个作者单位
    p4 = doc.add_paragraph("2. 清华大学，北京  100084")
    b4 = ParagraphBlock(index=3, paragraph=p4)
    b4.add_class("author-section")
    b4.add_class("author-affiliation")
    blocks.append(b4)
    
    # 5. 通讯作者
    p5 = doc.add_paragraph("*Email: zhangsan@example.com")
    b5 = ParagraphBlock(index=4, paragraph=p5)
    b5.add_class("author-section")
    b5.add_class("corresponding-author")
    blocks.append(b5)
    
    # 6. 摘要
    p6 = doc.add_paragraph("摘要：这是摘要内容...")
    b6 = ParagraphBlock(index=5, paragraph=p6)
    b6.add_class("abstract")
    blocks.append(b6)
    
    # 7. 关键词
    p7 = doc.add_paragraph("关键词：机器学习；深度学习；神经网络")
    b7 = ParagraphBlock(index=6, paragraph=p7)
    b7.add_class("keywords")
    blocks.append(b7)
    
    # 8. 引言标题
    p8 = doc.add_paragraph("引  言")
    b8 = ParagraphBlock(index=7, paragraph=p8)
    b8.add_class("heading-introduction")
    blocks.append(b8)
    
    # 9. 引言内容
    p9 = doc.add_paragraph("这是引言的第一段...")
    b9 = ParagraphBlock(index=8, paragraph=p9)
    b9.add_class("body-introduction")
    blocks.append(b9)
    
    # 10. 参考文献标题
    p10 = doc.add_paragraph("参考文献")
    b10 = ParagraphBlock(index=9, paragraph=p10)
    b10.add_class("heading-references")
    blocks.append(b10)
    
    # 11-13. 参考文献条目
    for i in range(1, 4):
        p = doc.add_paragraph(f"[{i}] 作者{i}. 论文标题{i}[J]. 期刊名, 2023.")
        b = ParagraphBlock(index=9+i, paragraph=p)
        b.add_class("reference-item")
        blocks.append(b)
    
    return blocks


def test_class_selector():
    """测试类选择器"""
    print("=" * 60)
    print("测试1: 类选择器")
    print("=" * 60)
    
    blocks = create_test_blocks()
    selector = Selector(blocks)
    
    # 测试选择作者单位
    results = selector.select(".author-affiliation")
    print(f"\n选择器: .author-affiliation")
    print(f"匹配数量: {len(results)}")
    for block in results:
        print(f"  - {block.paragraph.text}")
    
    assert len(results) == 2, "应该有2个作者单位"
    
    # 测试选择参考文献
    results = selector.select(".reference-item")
    print(f"\n选择器: .reference-item")
    print(f"匹配数量: {len(results)}")
    for block in results:
        print(f"  - {block.paragraph.text}")
    
    assert len(results) == 3, "应该有3条参考文献"
    
    print("\n✅ 类选择器测试通过！\n")


def test_pseudo_selector():
    """测试伪类选择器"""
    print("=" * 60)
    print("测试2: 伪类选择器")
    print("=" * 60)
    
    blocks = create_test_blocks()
    selector = Selector(blocks)
    
    # 测试 :first
    result = selector.select_one(".author-affiliation:first")
    print(f"\n选择器: .author-affiliation:first")
    print(f"结果: {result.paragraph.text if result else 'None'}")
    assert result and "北京大学" in result.paragraph.text, "应该是第一个作者单位"
    
    # 测试 :last
    result = selector.select_one(".author-affiliation:last")
    print(f"\n选择器: .author-affiliation:last")
    print(f"结果: {result.paragraph.text if result else 'None'}")
    assert result and "清华大学" in result.paragraph.text, "应该是第二个作者单位"
    
    # 测试 :nth(1)
    result = selector.select_one(".author-affiliation:nth(1)")
    print(f"\n选择器: .author-affiliation:nth(1)")
    print(f"结果: {result.paragraph.text if result else 'None'}")
    assert result and "清华大学" in result.paragraph.text, "应该是第二个作者单位（索引从0开始）"
    
    # 测试参考文献的第二条
    result = selector.select_one(".reference-item:nth(1)")
    print(f"\n选择器: .reference-item:nth(1)")
    print(f"结果: {result.paragraph.text if result else 'None'}")
    assert result and "[2]" in result.paragraph.text, "应该是第二条参考文献"
    
    print("\n✅ 伪类选择器测试通过！\n")


def test_adjacent_selector():
    """测试相邻兄弟选择器"""
    print("=" * 60)
    print("测试3: 相邻兄弟选择器")
    print("=" * 60)
    
    blocks = create_test_blocks()
    selector = Selector(blocks)
    
    # 测试选择引言标题后的引言内容
    result = selector.select_one(".heading-introduction + .body-introduction")
    print(f"\n选择器: .heading-introduction + .body-introduction")
    print(f"结果: {result.paragraph.text if result else 'None'}")
    assert result and "引言的第一段" in result.paragraph.text, "应该是引言内容"
    
    # 测试选择关键词后的引言标题
    result = selector.select_one(".keywords + .heading-introduction")
    print(f"\n选择器: .keywords + .heading-introduction")
    print(f"结果: {result.paragraph.text if result else 'None'}")
    assert result and "引  言" in result.paragraph.text, "应该是引言标题"
    
    print("\n✅ 相邻兄弟选择器测试通过！\n")


def test_utility_methods():
    """测试工具方法"""
    print("=" * 60)
    print("测试4: 工具方法")
    print("=" * 60)
    
    blocks = create_test_blocks()
    selector = Selector(blocks)
    
    # 测试 exists
    exists = selector.exists(".author-affiliation")
    print(f"\nexists('.author-affiliation'): {exists}")
    assert exists, "应该存在作者单位"
    
    exists = selector.exists(".non-existent")
    print(f"exists('.non-existent'): {exists}")
    assert not exists, "不应该存在不存在的类"
    
    # 测试 count
    count = selector.count(".reference-item")
    print(f"\ncount('.reference-item'): {count}")
    assert count == 3, "应该有3条参考文献"
    
    count = selector.count(".author-affiliation")
    print(f"count('.author-affiliation'): {count}")
    assert count == 2, "应该有2个作者单位"
    
    print("\n✅ 工具方法测试通过！\n")


def test_practical_examples():
    """测试实际应用示例"""
    print("=" * 60)
    print("测试5: 实际应用示例")
    print("=" * 60)
    
    blocks = create_test_blocks()
    selector = Selector(blocks)
    
    # 示例1: 查询第二个作者的地址
    print("\n示例1: 查询第二个作者的地址")
    result = selector.select_one(".author-affiliation:nth(1)")
    if result:
        print(f"第二个作者地址: {result.paragraph.text}")
        assert "清华大学" in result.paragraph.text
    
    # 示例2: 查询参考文献的第二条
    print("\n示例2: 查询参考文献的第二条")
    result = selector.select_one(".reference-item:nth(1)")
    if result:
        print(f"第二条参考文献: {result.paragraph.text}")
        assert "[2]" in result.paragraph.text
    
    # 示例3: 查询所有作者相关信息
    print("\n示例3: 查询所有作者相关信息")
    results = selector.select(".author-section")
    print(f"作者区域元素数量: {len(results)}")
    for block in results:
        print(f"  - {block.paragraph.text}")
    
    # 示例4: 查询引言标题后的第一段正文
    print("\n示例4: 查询引言标题后的第一段正文")
    result = selector.select_one(".heading-introduction + .body-introduction")
    if result:
        print(f"引言第一段: {result.paragraph.text}")
    
    print("\n✅ 实际应用示例测试通过！\n")


if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("开始 Selector 测试")
    print("=" * 60 + "\n")
    
    test_class_selector()
    test_pseudo_selector()
    test_adjacent_selector()
    test_utility_methods()
    test_practical_examples()
    
    print("=" * 60)
    print("✅ 所有测试通过！")
    print("=" * 60 + "\n")
