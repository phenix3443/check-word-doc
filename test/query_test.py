#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建一个用于测试 Selector 查询的文档
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 1. 标题
title = doc.add_paragraph("数据库（集）基本信息简介")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = "黑体"
    run.font.size = Pt(16)
    run.font.bold = True

# 2. 中文作者列表（多个作者）
authors = doc.add_paragraph("张三1，李四2，王五3")
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in authors.runs:
    run.font.name = "华文楷体"
    run.font.size = Pt(12)

# 3. 作者单位1
affiliation1 = doc.add_paragraph("1. 北京大学/计算机学院，北京  100871")
affiliation1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in affiliation1.runs:
    run.font.size = Pt(10.5)

# 4. 作者单位2
affiliation2 = doc.add_paragraph("2. 清华大学/软件学院，北京  100084")
affiliation2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in affiliation2.runs:
    run.font.size = Pt(10.5)

# 5. 作者单位3
affiliation3 = doc.add_paragraph("3. 中国科学院/计算技术研究所，北京  100190")
affiliation3.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in affiliation3.runs:
    run.font.size = Pt(10.5)

# 6. 通讯作者
corresponding = doc.add_paragraph("* 论文通信作者：张三（zhangsan@pku.edu.cn）")
corresponding.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in corresponding.runs:
    run.font.size = Pt(10.5)

# 7. 中文摘要
abstract = doc.add_paragraph("摘要：这是一个测试文档，用于演示 Selector 查询功能。本文介绍了如何使用类似 CSS 选择器的语法来查询文档中的特定元素。")
abstract.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in abstract.runs:
    run.font.name = "华文楷体"
    run.font.size = Pt(10.5)

# 8. 中文关键词
keywords = doc.add_paragraph("关键词：文档查询；选择器；CSS语法")
keywords.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in keywords.runs:
    run.font.name = "华文楷体"
    run.font.size = Pt(10.5)

# 9. 英文标题
title_en = doc.add_paragraph("Basic Information of Database")
title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title_en.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True

# 10. 英文作者列表
authors_en = doc.add_paragraph("Zhang San, Li Si, Wang Wu")
authors_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in authors_en.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

# 11. 英文通讯作者
corresponding_en = doc.add_paragraph("*Email: zhangsan@pku.edu.cn")
corresponding_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in corresponding_en.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)

# 12. 英文摘要
abstract_en = doc.add_paragraph("Abstract: This is a test document for demonstrating Selector query functionality.")
abstract_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in abstract_en.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)

# 13. 英文关键词
keywords_en = doc.add_paragraph("Keywords: Document Query; Selector; CSS Syntax")
keywords_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in keywords_en.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)

# 14. 一级标题：引言
heading1 = doc.add_paragraph("1 引言")
heading1.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in heading1.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)
    run.font.bold = True

# 15. 正文段落
body1 = doc.add_paragraph("这是引言部分的内容。介绍了研究背景和意义。")
body1.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in body1.runs:
    run.font.name = "宋体"
    run.font.size = Pt(10.5)

# 16. 一级标题：参考文献
heading_ref = doc.add_paragraph("参考文献")
heading_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in heading_ref.runs:
    run.font.name = "黑体"
    run.font.size = Pt(14)
    run.font.bold = True

# 17. 参考文献1
ref1 = doc.add_paragraph("[1]  张三, 李四. 文档处理技术研究[J]. 计算机学报, 2023, 46(1): 1-10.")
ref1.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in ref1.runs:
    run.font.name = "宋体"
    run.font.size = Pt(10.5)

# 18. 参考文献2
ref2 = doc.add_paragraph("[2]  Wang W, Li S. Document Analysis System[C]//Proceedings of ACL, 2023: 100-110.")
ref2.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in ref2.runs:
    run.font.name = "宋体"
    run.font.size = Pt(10.5)

# 19. 参考文献3
ref3 = doc.add_paragraph("[3]  Smith J. CSS Selectors Guide[M]. O'Reilly Media, 2022.")
ref3.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in ref3.runs:
    run.font.name = "宋体"
    run.font.size = Pt(10.5)

# 保存文档
output_path = "/Users/lsl/github/phenix3443/check-word-doc/test/query_test.docx"
doc.save(output_path)
print(f"✅ 测试文档已创建: {output_path}")
